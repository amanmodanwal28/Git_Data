from docx import Document
from datetime import datetime

# Create a new Document
doc = Document()

# Add a title
doc.add_heading("Project Documentation", 0)

# Add author and date
# doc.add_paragraph("Author: Your Name")
# doc.add_paragraph(f'Date: {datetime.now().strftime("%d-%m-%Y")}')

# Add a page break
doc.add_page_break()

# Add table of contents
doc.add_heading("Table of Contents", level=1)

toc = [
    "1. Introduction",
    "2. Project Overview",
    "3. Features and Functionalities",
    "    3.1 File Menu Options",
    "    3.2 Database Button",
    "    3.3 IP Address Management",
    "    3.4 Folder and File Management",
    "    3.5  File Addition Process",
    "4. Upload Process",
    "5. Error Handling and Validation",
    "6. Conclusion",
]

for item in toc:
    doc.add_paragraph(item, style="List Number")

# Add a page break
doc.add_page_break()

# Add Introduction section
doc.add_heading("1. Introduction", level=1)
doc.add_paragraph(
    "This document provides comprehensive documentation for the PPS project. "
    "The PPS project is designed to manage and monitor IP addresses, handle files and directories, and facilitate file uploads to selected IP addresses. "
    "The project features an intuitive graphical user interface (GUI) that allows users to interact with the system efficiently. "
    "This documentation covers all aspects of the project, including its features, functionalities, and the underlying processes."
)

doc.add_paragraph(
    "The project uses a configuration file (config.xml) to list IP addresses and determine their active or inactive status. "
    "The GUI refreshes every 15 seconds to update the status of IP addresses. Additionally, the project includes options for creating, opening, and managing projects, as well as uploading files to selected IP addresses."
)

# Add Project Overview section
doc.add_heading("2. Project Overview", level=1)
doc.add_paragraph(
    "The PPS project is a comprehensive system designed to manage and monitor IP addresses, handle file operations, and facilitate file uploads. "
    "The main components of the project include the following:\n"
    "- **IP Address Management**: Lists IP addresses from the config.xml file and indicates their active or inactive status.\n"
    "- **Folder and File Management**: Allows users to create, open, and manage project directories, add and delete files, and perform batch operations on files.\n"
    "- **Upload Process**: Facilitates the uploading of files to selected IP addresses, ensuring data integrity through CRC value validation.\n"
)

doc.add_paragraph(
    "The overall workflow of the project involves the following steps:\n"
    "1. **IP Address Monitoring**: The system reads IP addresses from the config.xml file and displays their status in the GUI.\n"
    "2. **Project Management**: Users can create new projects, open existing ones, and manage directories and files within the projects.\n"
    "3. **File Operations**: Users can add files from their local system to the project directory, delete individual files or all files, and generate an INDEX.SYS file that contains the file paths, sizes, and CRC values.\n"
    "4. **File Upload**: Users can select multiple IP addresses and upload files to them. The system validates the CRC values before uploading to ensure data integrity.\n"
)

# Add Features and Functionalities section
doc.add_heading("3. Features and Functionalities", level=1)

# File Menu Options
doc.add_heading("3.1 File Menu Options", level=2)
doc.add_paragraph(
    "The file menu provides options for creating new projects, opening existing projects, and managing directories and files. Here are the details:\n"
    "- **New Project**: This option allows users to create a new project. When selected, a blank project is created, and the user is prompted to choose a project directory. "
    "After creating the project directory, necessary subfolders are added, and a blank INDEX.SYS file is created.\n"
    "- **Open Project**: This option allows users to open an existing project. When selected, the user can choose a project directory, and the contents of the directory are displayed in the GUI.\n"
    "If the selected project folder does not contain the necessary subfolders as specified in the config.xml file, a dialog is displayed informing the user that the project folder is incorrect.\n"
    "- **Database Option**: Similar to the New Project and Open Project options, this allows users to manage databases within the project."
)
doc.add_paragraph(
    "Example of creating a new project with necessary subfolders and INDEX.SYS file:\n"
    "1. The user selects 'New Project' from the file menu.\n"
    "2. The system prompts the user to choose a directory for the new project.\n"
    "3. The system creates the chosen directory and adds necessary subfolders (e.g., 'data', 'logs', 'config').\n"
    "4. A blank INDEX.SYS file is created in the project directory.\n"
)

# Database Button
doc.add_heading("3.4 Database Button", level=2)
doc.add_paragraph(
    "The main screen includes a 'Database' button that functions similarly to the 'Open Project' option. Here are the details:\n"
    "- **Database Option**: This button allows users to manage databases within the project. When selected, the user can choose a database directory, and the contents of the directory are displayed in the GUI. "
    "If the selected database folder does not contain the necessary subfolders as specified in the config.xml file, a dialog is displayed informing the user that the database folder is incorrect."
)
# IP Address Management
doc.add_heading("3.2 IP Address Management", level=2)
doc.add_paragraph(
    "The IP address management functionality is designed to monitor and display the status of IP addresses listed in the config.xml file. The key features include:\n"
    "- **IP List Display**: The GUI lists IP addresses retrieved from the config.xml file. Each IP address is displayed with a status indicator.\n"
    "- **Status Indicators**: The status of each IP address is updated every 15 seconds. A green checkmark (✔️) indicates an active IP, while a red cross (❌) indicates an inactive IP.\n"
    "- **Multiple Selection**: Users can select multiple IP addresses using checkboxes provided next to each IP."
)

# Folder and File Management
doc.add_heading("3.3 Folder and File Management", level=2)
doc.add_paragraph(
    "The folder and file management functionality allows users to handle project directories and files efficiently. The main features include:\n"
    "- **Directory Display**: The selected project directory and its contents are displayed in the GUI.\n"
    "- **Add Files**: Users can add files from their local system to the project directory. When files are added, an INDEX.SYS file is generated, which includes the file paths, sizes, and CRC values.\n"
    "- **Delete Files**: Users can delete individual files or all files within the project directory. Before deleting, a confirmation prompt is displayed to ensure the user wants to proceed with the deletion.\n"
    "- **Batch Operations**: When files are added or deleted, batch operations can be performed. For example, users can select all files and perform a delete operation in one go."
)

doc.add_heading("3.4 File Addition Process", level=1)
doc.add_paragraph(
    "The file addition process involves adding files from the local system to the project folder and generating an INDEX.SYS file with details about the added files. The steps are as follows:\n"
    "- **File Selection**: Users can select files from their local system to add to the project folder.\n"
    "- **Data Copying**: A popup is displayed showing how much data is being copied to the project. If any files are not added (e.g., unsupported file formats like certain audio or video types), a dialog box notifies the user which files were not added.\n"
    "- **Progress Indication**: During the upload process, a progress bar is displayed to indicate the progress of the file transfer, providing real-time feedback on the upload status.\n"
    "- **Format Conversion**: When adding files, certain formats are converted to ensure compatibility. For instance, all audio files are converted to MP3 format except for specific formats like M4V, and all video files are converted to MP4 format except for certain formats like ('webm','wmv','ogv','flv') ,except for certain other files formats like('alac', 'svg', 'webp', 'psd', 'bmp', 'ico').\n"
    "- **INDEX.SYS Generation**: As files are added, an INDEX.SYS file is generated, containing the paths, sizes, and CRC values of the added files. This file is used during the upload process to validate the files."
)

# Add Upload Process section
doc.add_heading("4. Upload Process", level=1)
doc.add_paragraph(
    "The upload process involves selecting IP addresses, validating the files, and uploading the entire project folder to the selected IP addresses. The steps are as follows:\n"
    "- **IP Selection**: Users can select multiple IP addresses to upload the project to. If no IP addresses are selected, a popup is shown instructing the user to select IP addresses before proceeding.\n"
    "- **CRC Validation**: Before uploading, the system generates CRC values for each file in the project directory and compares them with the values in the INDEX.SYS file. If the values match, the files are uploaded; otherwise, an error popup with error code 12345 is displayed, and the files with mismatched CRC values are not copied.\n"
    "- **Upload Confirmation**: Once the IP addresses are selected and the CRC validation is complete, users are prompted to confirm the upload process, specifying which IP addresses to send the files to.\n"
    "- **Full Project Upload**: Upon confirmation, the system uploads the entire project directory as per the INDEX.SYS file to the selected IP addresses.\n"
    "- **Progress Indication**: During the upload process, a progress bar is displayed to indicate the progress of the file transfer, providing real-time feedback on the upload status.\n"
    "- **Error Handling**: If there is a connection error during the upload, a dialog popup is displayed informing the user of the connection loss and that the files were not copied."
)


# Add Error Handling and Validation section
doc.add_heading("5. Error Handling and Validation", level=1)
doc.add_paragraph(
    "The system includes robust error handling and validation mechanisms to ensure data integrity and user experience. Key aspects include:\n"
    "- **CRC Validation**: Ensures that the files being uploaded match the CRC values in the INDEX.SYS file, preventing data corruption. Files with mismatched CRC values are not uploaded, and the user is notified with a specific error code (e.g., 12345).\n"
    "- **Confirmation Prompts**: Provides confirmation dialogs for critical actions like deleting files or uploading data, preventing accidental operations. For example, before deleting files, the system asks for user confirmation.\n"
    "- **Error Popups**: Displays error messages when validation fails, guiding users to take corrective actions. For instance, if there is a connection error during the upload, the system shows a popup indicating the connection loss and which files were not copied.\n"
    "- **Project Folder Validation**: When the user selects a project folder, the system checks if the folder includes the necessary subfolders as specified in the config.xml file. If the selected folder does not contain the required subfolders, a dialog is displayed informing the user that the project folder is incorrect."
)

# Add Conclusion section
doc.add_heading("6. Conclusion", level=1)
doc.add_paragraph(
    "The PPS project provides a comprehensive solution for managing and monitoring IP addresses, handling project directories and files, and facilitating secure file uploads. Its intuitive GUI and robust functionality make it a valuable tool for efficient project management."
)

# Save the final document
doc.save("Project_Documentation_Complete.docx")
